"""
Retirement Financial Modeling — Report Generator.

Produces a professional .docx report (and optional .pdf conversion) from
a completed SimulationResults object. Called on demand from the Streamlit app.

Report structure:
  1. Cover page — title, run date, key headline metrics
  2. Executive Summary — plain-language interpretation of results
  3. Charts — portfolio fan, income stack, tax burden, asset allocation (PNG, embedded)
  4. Data Tables — year-by-year percentile table, tax schedule, SS income schedule
  5. Assumptions Log — full SimulationConfig documented for auditor/CPA traceability

Dependencies:
  - python-docx >= 1.1.0
  - matplotlib >= 3.7
  - numpy, pandas (already in requirements)
  - LibreOffice (for PDF conversion, optional)

Output:
  outputs/retirement_report_YYYYMMDD_HHMMSS.docx
  outputs/retirement_report_YYYYMMDD_HHMMSS.pdf  (if LibreOffice available)

Version: 1.7.0
"""

import os
import io
import math
import re
import zipfile
import subprocess
import tempfile
import logging
from datetime import datetime
from typing import Optional

import numpy as np
import matplotlib
matplotlib.use("Agg")  # Non-interactive backend for server-side rendering
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from config.defaults import SimulationConfig, SimulationResults, ASSET_CLASSES

logger = logging.getLogger(__name__)

# -- Visual constants ---------------------------------------------------------
CHART_BG      = "#0A0A0A"   # near-black chart background
CHART_PLOT_BG = "#111111"   # slightly lighter plot area
CHART_GRID    = "#2A2A2A"   # subtle grid lines
CHART_TEXT    = "#FFFFFF"   # axis labels, ticks
ACCENT_BLUE   = "#4A9EFF"   # primary accent
ACCENT_GREEN  = "#00CC88"   # success / positive
ACCENT_RED    = "#FF4444"   # failure / negative
ACCENT_YELLOW = "#FFD700"   # warning / neutral

# Percentile band colors
P10_COLOR  = "#FF4444"
P25_COLOR  = "#FF8844"
P50_COLOR  = "#FFFFFF"
P75_COLOR  = "#88CC44"
P90_COLOR  = "#00CC88"

# -- Formatting helpers -------------------------------------------------------

def _fmt_dollar(v: float) -> str:
    """Unified monetary formatter — auto-selects K or M by magnitude."""
    if math.isnan(v):
        return "$--"
    if math.isinf(v):
        return "$--"
    abs_v = abs(v)
    sign = "-" if v < 0 else ""
    if abs_v >= 1_000_000:
        return f"{sign}${abs_v / 1_000_000:,.2f}M"
    elif abs_v >= 1_000:
        return f"{sign}${abs_v / 1_000:,.1f}K"
    return f"{sign}${abs_v:,.0f}"


_fmt_k = _fmt_dollar
_fmt_m = _fmt_dollar


def _pct(v: float, decimals: int = 1) -> str:
    """Format fraction as percentage: 0.873 -> 87.3%."""
    return f"{v * 100:.{decimals}f}%"


def _age_label(ages: np.ndarray) -> list:
    """Return every 5th age as a string label for chart x-axis."""
    return [str(int(a)) if int(a) % 5 == 0 else "" for a in ages]


# -- Chart generators (matplotlib, dark theme) --------------------------------

def _apply_dark_style(fig: plt.Figure, ax) -> None:
    """Apply consistent dark theme to a matplotlib figure/axis."""
    fig.patch.set_facecolor(CHART_BG)
    ax.set_facecolor(CHART_PLOT_BG)
    ax.tick_params(colors=CHART_TEXT, labelsize=9)
    ax.xaxis.label.set_color(CHART_TEXT)
    ax.yaxis.label.set_color(CHART_TEXT)
    ax.title.set_color(CHART_TEXT)
    for spine in ax.spines.values():
        spine.set_edgecolor(CHART_GRID)
    ax.grid(True, color=CHART_GRID, linewidth=0.5, linestyle="--", alpha=0.7)


def _chart_portfolio_fan(results: SimulationResults) -> io.BytesIO:
    """
    Portfolio fan chart — percentile bands across all simulation years.

    Args:
        results: Completed SimulationResults object.

    Returns:
        BytesIO PNG image buffer.
    """
    ages = results.ages
    pv   = results.portfolio_values  # (n_sims, n_years)

    p10 = np.percentile(pv, 10, axis=0) / 1e6
    p25 = np.percentile(pv, 25, axis=0) / 1e6
    p50 = np.percentile(pv, 50, axis=0) / 1e6
    p75 = np.percentile(pv, 75, axis=0) / 1e6
    p90 = np.percentile(pv, 90, axis=0) / 1e6

    fig, ax = plt.subplots(figsize=(9, 4.5), dpi=150)
    _apply_dark_style(fig, ax)

    ax.fill_between(ages, p10, p90, alpha=0.15, color=ACCENT_BLUE, label="_nolegend_")
    ax.fill_between(ages, p25, p75, alpha=0.25, color=ACCENT_BLUE, label="_nolegend_")
    ax.plot(ages, p10, color=P10_COLOR,  linewidth=1.0, linestyle="--", label="10th pctile")
    ax.plot(ages, p25, color=P25_COLOR,  linewidth=1.0, linestyle="--", label="25th pctile")
    ax.plot(ages, p50, color=P50_COLOR,  linewidth=2.0,                 label="Median (50th)")
    ax.plot(ages, p75, color=P75_COLOR,  linewidth=1.0, linestyle="--", label="75th pctile")
    ax.plot(ages, p90, color=P90_COLOR,  linewidth=1.0, linestyle="--", label="90th pctile")

    # Retirement age marker
    ret_age = results.config.retirement_age
    ax.axvline(x=ret_age, color=ACCENT_YELLOW, linewidth=1.2, linestyle=":", alpha=0.8)
    ax.text(ret_age + 0.3, ax.get_ylim()[1] * 0.95 if ax.get_ylim()[1] > 0 else 1,
            "Retirement", color=ACCENT_YELLOW, fontsize=8, va="top")

    ax.set_xlabel("Age", fontsize=10)
    ax.set_ylabel("Portfolio Value ($M)", fontsize=10)
    ax.set_title("Portfolio Value — Monte Carlo Percentile Bands", fontsize=12, pad=10)
    ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, _: f"${x:.1f}M"))

    legend = ax.legend(fontsize=8, facecolor=CHART_BG, edgecolor=CHART_GRID,
                       labelcolor=CHART_TEXT, loc="upper left")
    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", facecolor=CHART_BG, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _chart_income_stack(results: SimulationResults) -> io.BytesIO:
    """
    Stacked area chart of income sources across retirement (median path).

    Args:
        results: Completed SimulationResults object.

    Returns:
        BytesIO PNG image buffer.
    """
    ages   = results.ages
    config = results.config

    # Median withdrawals and SS
    median_withdrawal = np.median(results.withdrawals, axis=0)
    median_ss         = np.median(results.ss_income,   axis=0)

    # Build other income series (deterministic -- same every path)
    pension   = np.array([
        config.pension_annual_real
        if config.pension_start_age <= age <= config.pension_end_age else 0.0
        for age in ages
    ])
    rental    = np.array([
        config.rental_income_annual_real
        if config.rental_start_age <= age <= config.rental_end_age else 0.0
        for age in ages
    ])
    part_time = np.array([
        config.part_time_income_annual
        if config.part_time_income_start_age <= age <= config.part_time_income_end_age else 0.0
        for age in ages
    ])
    # SERP per-year schedule (nominal contractual amounts, 2026-2033)
    _start_yr = getattr(config, 'simulation_start_year', 2026)
    serp = np.array([
        getattr(config, f'serp_{_start_yr + (age - config.current_age)}', 0.0)
        for age in ages
    ])

    # Portfolio withdrawal = total withdrawal minus other income sources
    other_total = median_ss + pension + rental + part_time + serp
    portfolio_w = np.maximum(median_withdrawal - other_total, 0)

    fig, ax = plt.subplots(figsize=(9, 4.5), dpi=150)
    _apply_dark_style(fig, ax)

    # Only include series with non-zero values to keep chart clean
    labels_all = ["Portfolio", "Social Security", "Pension", "Rental",
                  "Part-Time", "SERP"]
    data_all   = [portfolio_w, median_ss, pension, rental, part_time, serp]
    colors_all = [ACCENT_BLUE, ACCENT_GREEN, "#FF8844", "#CC44FF",
                  "#44CCFF", "#FFD700"]
    active = [(l, d, c) for l, d, c in zip(labels_all, data_all, colors_all)
              if np.any(d > 0)]
    labels, data, colors = zip(*active) if active else (labels_all[:2], data_all[:2], colors_all[:2])

    ax.stackplot(ages, data, labels=labels, colors=colors, alpha=0.85)
    ax.set_xlabel("Age", fontsize=10)
    ax.set_ylabel("Annual Income ($)", fontsize=10)
    ax.set_title("Income Sources — Median Simulation Path", fontsize=12, pad=10)
    ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, _: f"${x:,.0f}"))
    legend = ax.legend(fontsize=8, facecolor=CHART_BG, edgecolor=CHART_GRID,
                       labelcolor=CHART_TEXT, loc="upper right")
    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", facecolor=CHART_BG, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _chart_tax_burden(results: SimulationResults) -> io.BytesIO:
    """
    Annual tax burden (median) and effective rate across retirement.

    Args:
        results: Completed SimulationResults object.

    Returns:
        BytesIO PNG image buffer.
    """
    ages        = results.ages
    median_tax  = np.median(results.taxes, axis=0)
    if hasattr(results, 'gross_income') and results.gross_income.size > 0:
        median_gross = np.median(results.gross_income, axis=0)
    else:
        median_gross = np.zeros_like(median_tax)
    with np.errstate(divide='ignore', invalid='ignore'):
        eff_rate = np.where(median_gross > 0, median_tax / median_gross, 0.0) * 100

    fig, ax1 = plt.subplots(figsize=(9, 4.0), dpi=150)
    _apply_dark_style(fig, ax1)

    ax1.bar(ages, median_tax / 1000, color=ACCENT_BLUE, alpha=0.75, label="Annual Tax ($K)")
    ax1.set_xlabel("Age", fontsize=10)
    ax1.set_ylabel("Annual Tax ($K)", fontsize=10, color=ACCENT_BLUE)
    ax1.tick_params(axis="y", labelcolor=ACCENT_BLUE)

    ax2 = ax1.twinx()
    ax2.set_facecolor(CHART_PLOT_BG)
    ax2.plot(ages, eff_rate, color=ACCENT_YELLOW, linewidth=2.0, label="Effective Rate (%)")
    ax2.set_ylabel("Effective Tax Rate (%)", fontsize=10, color=ACCENT_YELLOW)
    ax2.tick_params(axis="y", labelcolor=ACCENT_YELLOW, colors=CHART_TEXT)
    for spine in ax2.spines.values():
        spine.set_edgecolor(CHART_GRID)

    ax1.set_title("Annual Tax Burden — Median Path", fontsize=12, pad=10)
    lines1, labels1 = ax1.get_legend_handles_labels()
    lines2, labels2 = ax2.get_legend_handles_labels()
    ax1.legend(lines1 + lines2, labels1 + labels2, fontsize=8,
               facecolor=CHART_BG, edgecolor=CHART_GRID, labelcolor=CHART_TEXT)
    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", facecolor=CHART_BG, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _chart_withdrawal_schedule(results: SimulationResults) -> io.BytesIO:
    """
    Median household withdrawal schedule chart.

    Args:
        results: Completed SimulationResults object.

    Returns:
        BytesIO PNG image buffer.
    """
    ages = results.ages
    median_wd = np.median(results.withdrawals, axis=0) / 1e3

    fig, ax = plt.subplots(figsize=(9, 4.5), dpi=150)
    _apply_dark_style(fig, ax)

    ax.fill_between(ages, 0, median_wd, color=ACCENT_BLUE, alpha=0.75, label="Household Withdrawal")

    # Retirement age marker
    ret_age = results.config.retirement_age
    ax.axvline(x=ret_age, color=ACCENT_YELLOW, linewidth=1.2, linestyle=":", alpha=0.8)
    ax.text(ret_age + 0.3, ax.get_ylim()[1] * 0.95 if ax.get_ylim()[1] > 0 else 1,
            "Retirement", color=ACCENT_YELLOW, fontsize=8, va="top")

    ax.set_xlabel("Age", fontsize=10)
    ax.set_ylabel("Annual Withdrawal ($K)", fontsize=10)
    ax.set_title("Household Withdrawal Schedule — Median Path", fontsize=12, pad=10)
    ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, _: f"${x:,.0f}K"))
    ax.legend(fontsize=8, facecolor=CHART_BG, edgecolor=CHART_GRID,
              labelcolor=CHART_TEXT, loc="upper left")
    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", facecolor=CHART_BG, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _chart_allocation_pies(config: SimulationConfig) -> io.BytesIO:
    """
    Side-by-side pie charts of pre- and post-retirement asset allocations.

    Args:
        config: SimulationConfig with allocation dicts.

    Returns:
        BytesIO PNG image buffer.
    """
    labels = [a.replace("_", " ").title() for a in ASSET_CLASSES]
    colors = [
        "#4A9EFF", "#00CC88", "#FF8844", "#CC44FF",
        "#FFD700", "#44CCFF", "#FF4444", "#88CC44"
    ]

    pre_vals = [config.pre_retirement_allocation.get(a, 0) for a in ASSET_CLASSES]
    ret_vals = [config.retirement_allocation.get(a, 0) for a in ASSET_CLASSES]

    # Filter zero slices
    def nonzero(vals):
        return [(l, v, c) for l, v, c in zip(labels, vals, colors) if v > 0]

    pre_nz  = nonzero(pre_vals)
    ret_nz  = nonzero(ret_vals)

    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(9, 4.0), dpi=150)
    fig.patch.set_facecolor(CHART_BG)

    for ax, data, title in [
        (ax1, pre_nz, "Pre-Retirement"),
        (ax2, ret_nz, "Retirement"),
    ]:
        ax.set_facecolor(CHART_BG)
        if data:
            ls, vs, cs = zip(*data)
            wedges, texts, autotexts = ax.pie(
                vs, labels=ls, colors=cs, autopct="%1.0f%%",
                startangle=90, pctdistance=0.75,
                textprops={"color": CHART_TEXT, "fontsize": 8},
                wedgeprops={"linewidth": 0.5, "edgecolor": CHART_BG},
            )
            for at in autotexts:
                at.set_color(CHART_BG)
                at.set_fontsize(7)
        ax.set_title(title, color=CHART_TEXT, fontsize=11, pad=6)

    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", facecolor=CHART_BG, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _chart_success_curve(results: SimulationResults) -> io.BytesIO:
    """
    Portfolio success curve — year-by-year survival probability.

    Args:
        results: Completed SimulationResults object.

    Returns:
        BytesIO PNG image buffer.
    """
    n_sims, n_years = results.portfolio_values.shape
    survival_pct = np.zeros(n_years)
    for t in range(n_years):
        alive = np.sum(results.portfolio_values[:, t] > 0)
        survival_pct[t] = (alive / n_sims) * 100.0

    years = np.arange(1, n_years + 1)

    fig, ax = plt.subplots(figsize=(9, 3.5), dpi=150)
    _apply_dark_style(fig, ax)

    ax.fill_between(years, 0, survival_pct, alpha=0.35, color=ACCENT_BLUE)
    ax.plot(years, survival_pct, color=ACCENT_BLUE, linewidth=2)

    ax.set_xlabel("Year", fontsize=10)
    ax.set_ylabel("Success (%)", fontsize=10)
    ax.set_title("Portfolio Success — Year-by-Year Survival Probability", fontsize=12, pad=10)
    ax.set_ylim(0, 105)
    ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, _: f"{x:.0f}%"))

    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", facecolor=CHART_BG, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _chart_balance_lines(results: SimulationResults) -> io.BytesIO:
    """
    Simulated portfolio balances — 5 percentile lines.

    Args:
        results: Completed SimulationResults object.

    Returns:
        BytesIO PNG image buffer.
    """
    pv = results.portfolio_values
    n_sims, n_years = pv.shape
    years = np.arange(0, n_years)

    p10 = np.percentile(pv, 10, axis=0)
    p25 = np.percentile(pv, 25, axis=0)
    p50 = np.percentile(pv, 50, axis=0)
    p75 = np.percentile(pv, 75, axis=0)
    p90 = np.percentile(pv, 90, axis=0)

    fig, ax = plt.subplots(figsize=(9, 4.5), dpi=150)
    _apply_dark_style(fig, ax)

    ax.plot(years, p10, color="#E63946", linewidth=1.5, label="10th Percentile")
    ax.plot(years, p25, color="#F4A261", linewidth=1.5, label="25th Percentile")
    ax.plot(years, p50, color="#E9C46A", linewidth=2.0, label="50th Percentile")
    ax.plot(years, p75, color="#2A9D8F", linewidth=1.5, label="75th Percentile")
    ax.plot(years, p90, color="#457B9D", linewidth=1.5, label="90th Percentile")

    ax.set_xlabel("Year", fontsize=10)
    ax.set_ylabel("Portfolio Balance ($)", fontsize=10)
    ax.set_title("Simulated Portfolio Balances", fontsize=12, pad=10)
    ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, _: f"${x/1e6:.1f}M" if x >= 1e6 else f"${x/1e3:.0f}K"))

    legend = ax.legend(fontsize=8, facecolor=CHART_BG, edgecolor=CHART_GRID,
                       labelcolor=CHART_TEXT, loc="upper left")
    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", facecolor=CHART_BG, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


# -- python-docx helpers ------------------------------------------------------

def _set_cell_bg(cell, hex_color: str) -> None:
    """Set table cell background color via XML shading."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.lstrip("#"))
    tcPr.append(shd)


def _heading(doc: Document, text: str, level: int) -> None:
    """Add a heading paragraph."""
    p = doc.add_heading(text, level=level)
    p.runs[0].font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C) if level == 1 else RGBColor(0x2E, 0x5F, 0x8A)


def _body(doc: Document, text: str) -> None:
    """Add a normal body paragraph."""
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(11)


def _add_table(doc: Document, headers: list, rows: list,
               col_widths: Optional[list] = None) -> None:
    """
    Add a formatted table with header row and alternating row shading.

    Args:
        doc: Document to add the table to.
        headers: List of column header strings.
        rows: List of lists of cell values (will be str-converted).
        col_widths: Optional list of column widths in inches.
    """
    n_cols = len(headers)
    table  = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold      = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_bg(cell, "#1A3A5C")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row_obj = table.rows[r_idx + 1]
        bg      = "F5F8FC" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row_obj.cells[c_idx]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            _set_cell_bg(cell, bg)

    # Column widths
    if col_widths:
        for row_obj in table.rows:
            for c_idx, width in enumerate(col_widths):
                row_obj.cells[c_idx].width = Inches(width)


# -- Report generator ---------------------------------------------------------

class ReportGenerator:
    """
    Generates a professional .docx (and optional .pdf) retirement analysis report.

    The report includes an executive summary, embedded charts, year-by-year
    data tables, and a complete assumptions log for auditor / CPA review.

    Attributes:
        results: SimulationResults from MonteCarloSimulator.run().
        config: SimulationConfig (shorthand for results.config).
        output_dir: Directory where reports are saved.
    """

    def __init__(self, results: SimulationResults,
                 output_dir: str = "outputs"):
        """
        Initialize the report generator.

        Args:
            results: Completed SimulationResults object.
            output_dir: Directory path for saving output files.
        """
        self.results    = results
        self.config     = results.config
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def _next_filename(self) -> str:
        """
        Generate the next versioned filename for today's date.

        Pattern: YYYY Month DD Retirement Report v#
        Example: 2026 March 09 Retirement Report v1

        Auto-increments version if a file with the same date already exists.

        Returns:
            Filename base string (no extension).
        """
        now    = datetime.now()
        prefix = now.strftime("%Y %B %d Retirement Report")
        version = 1
        while True:
            candidate = f"{prefix} v{version}"
            if not (os.path.exists(os.path.join(self.output_dir, candidate + ".docx")) or
                    os.path.exists(os.path.join(self.output_dir, candidate + ".pdf"))):
                return candidate
            version += 1

    def generate(self, filename_base: Optional[str] = None) -> dict:
        """
        Generate .docx and .pdf reports.

        Args:
            filename_base: Base filename without extension.
                           Defaults to retirement_report_YYYYMMDD_HHMMSS.

        Returns:
            Dict with keys 'docx' and 'pdf' mapping to absolute file paths.
            'pdf' is None if LibreOffice is not available.
        """
        if filename_base is None:
            filename_base = self._next_filename()

        docx_path = os.path.join(self.output_dir, f"{filename_base}.docx")

        logger.info("Generating charts...")
        charts = self._generate_charts()

        logger.info("Building Word document...")
        self._build_docx(docx_path, charts)
        self._patch_docx(docx_path)

        logger.info("Converting to PDF...")
        pdf_path = self._convert_to_pdf(docx_path)

        logger.info("Report complete: %s", docx_path)
        return {"docx": os.path.abspath(docx_path),
                "pdf":  os.path.abspath(pdf_path) if pdf_path else None}

    # -- Chart generation -----------------------------------------------------

    def _generate_charts(self) -> dict:
        """
        Render all report charts as PNG BytesIO buffers.

        Returns:
            Dict mapping chart name -> BytesIO buffer.
        """
        return {
            "portfolio_fan":        _chart_portfolio_fan(self.results),
            "income_stack":         _chart_income_stack(self.results),
            "tax_burden":           _chart_tax_burden(self.results),
            "allocation_pies":      _chart_allocation_pies(self.config),
            "withdrawal_schedule":  _chart_withdrawal_schedule(self.results),
            "success_curve":        _chart_success_curve(self.results),
            "balance_lines":        _chart_balance_lines(self.results),
        }

    # -- Document builder -----------------------------------------------------

    def _build_docx(self, path: str, charts: dict) -> None:
        """
        Build the complete Word document.

        Args:
            path: Output file path.
            charts: Dict of chart name -> PNG BytesIO from _generate_charts().
        """
        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin    = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin   = Cm(2.5)
            section.right_margin  = Cm(2.5)

        self._add_cover(doc)
        self._add_executive_summary(doc)
        self._add_charts_section(doc, charts)
        self._add_data_tables(doc)
        self._add_income_breakdown_table(doc)
        self._add_common_sense_check(doc)
        self._add_rich_statistics(doc)
        self._add_assumptions_log(doc)

        doc.save(path)

    def _add_cover(self, doc: Document) -> None:
        """Add cover page."""
        doc.add_paragraph()
        doc.add_paragraph()

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("Retirement Financial Analysis")
        run.bold      = True
        run.font.size = Pt(28)
        run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = subtitle.add_run("Monte Carlo Simulation Report")
        run2.font.size = Pt(16)
        run2.font.color.rgb = RGBColor(0x2E, 0x5F, 0x8A)

        doc.add_paragraph()

        meta_lines = [
            f"Generated: {datetime.now().strftime('%B %d, %Y at %H:%M')}",
            f"Simulations: {self.config.n_simulations:,} paths  |  "
            f"Method: {self.config.simulation_method.upper()}",
            f"Planning Horizon: Age {self.config.current_age} -> Age {self.config.life_expectancy}",
        ]
        for line in meta_lines:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line)
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        doc.add_page_break()

    def _add_executive_summary(self, doc: Document) -> None:
        """Add executive summary section with headline metrics and interpretation."""
        _heading(doc, "1. Executive Summary", 1)

        r  = self.results
        c  = self.config
        pv = r.portfolio_values

        final_vals  = pv[:, -1]
        success_pct = r.success_rate * 100
        median_final = np.median(final_vals)
        p10_final    = np.percentile(final_vals, 10)
        p90_final    = np.percentile(final_vals, 90)

        # Headline metric table
        median_tax     = np.mean(np.median(r.taxes,       axis=0))
        if hasattr(r, 'gross_income') and r.gross_income.size > 0:
            median_gross = np.mean(np.median(r.gross_income, axis=0))
        else:
            median_gross = 0.0
        eff_rate       = median_tax / median_gross if median_gross > 0 else 0

        _add_table(doc,
            headers=["Metric", "Value"],
            rows=[
                ["Success Rate (portfolio survives to life expectancy)",
                 f"{success_pct:.1f}%"],
                ["Median Portfolio at Life Expectancy (Age {})".format(c.life_expectancy),
                 _fmt_m(median_final)],
                ["10th Percentile Portfolio at Life Expectancy",
                 _fmt_m(p10_final)],
                ["90th Percentile Portfolio at Life Expectancy",
                 _fmt_m(p90_final)],
                ["Target Annual Withdrawal (Today's Dollars)",
                 _fmt_k(c.annual_withdrawal_real)],
                ["Median Annual Tax Paid",
                 _fmt_m(median_tax)],
                ["Median Effective Tax Rate",
                 _pct(eff_rate)],
                ["Simulations Run",
                 f"{c.n_simulations:,}"],
            ],
            col_widths=[4.5, 2.0]
        )

        doc.add_paragraph()
        _heading(doc, "Interpretation", 2)

        # Plain-language interpretation
        if success_pct >= 90:
            outlook = (
                f"The portfolio demonstrates strong resilience with a {success_pct:.1f}% success rate. "
                f"In {success_pct:.1f}% of simulated market environments -- spanning recessions, "
                f"inflationary periods, and bull markets drawn from {r.ages[0]} to {r.ages[-1]} "
                f"years of historical data -- the portfolio funded withdrawals through age "
                f"{c.life_expectancy} without depletion."
            )
        elif success_pct >= 75:
            outlook = (
                f"The portfolio shows moderate resilience with a {success_pct:.1f}% success rate. "
                f"While the majority of scenarios result in full funding through age {c.life_expectancy}, "
                f"approximately {100 - success_pct:.1f}% of simulated market sequences led to portfolio "
                f"depletion before that age. Consider reducing withdrawal targets or adjusting asset allocation."
            )
        else:
            outlook = (
                f"The portfolio faces material risk with a {success_pct:.1f}% success rate. "
                f"In approximately {100 - success_pct:.1f}% of simulated scenarios, the portfolio "
                f"depleted before age {c.life_expectancy}. Significant adjustments to withdrawal "
                f"rate, asset allocation, or retirement timing are recommended."
            )

        _body(doc, outlook)
        doc.add_paragraph()

        ss_annual = c.ss_monthly_benefit_at_fra * 12
        _body(doc,
            f"Social Security income of {_fmt_k(ss_annual)}/year (claimed at age "
            f"{c.ss_claiming_age}) provides a meaningful income floor. The three-bucket "
            f"withdrawal strategy (Taxable -> Traditional -> Roth) optimizes tax efficiency "
            f"by deferring tax-free Roth assets for as long as possible."
        )

        # SERP per-year schedule callout
        _serp_items = [(yr, getattr(c, f'serp_{yr}', 0.0)) for yr in range(2026, 2034)]
        _serp_total = sum(v for _, v in _serp_items)
        if _serp_total > 0:
            doc.add_paragraph()
            _heading(doc, "SERP / Nonqualified Deferred Compensation", 2)
            _sched_str = ", ".join(f"{yr}: {_fmt_k(v)}" for yr, v in _serp_items if v > 0)
            _body(doc,
                f"SERP distribution schedule (nominal dollars): {_sched_str}. "
                f"Total: {_fmt_k(_serp_total)} over {sum(1 for _, v in _serp_items if v > 0)} years. "
                f"Tax treatment: ordinary income for federal and state income tax in the "
                f"year of distribution (IRC 409A). FICA (Social Security + Medicare) does "
                f"NOT apply at distribution -- FICA was paid at the time of deferral when "
                f"the compensation was earned."
            )

        doc.add_page_break()

    def _add_charts_section(self, doc: Document, charts: dict) -> None:
        """Embed all charts as PNG images in the document."""
        _heading(doc, "2. Charts", 1)

        chart_meta = [
            ("portfolio_fan",   "Portfolio Value — Monte Carlo Percentile Bands",
             "Fan chart showing the 10th, 25th, 50th, 75th, and 90th percentile portfolio "
             "balance at each age. The shaded region represents the interquartile range (25th-75th). "
             "The dashed vertical line marks retirement age."),
            ("income_stack",    "Income Sources — Median Simulation Path",
             "Stacked area chart showing how total retirement income is composed across age. "
             "Sources include portfolio withdrawals, Social Security, pension, rental income, "
             "and part-time work income."),
            ("tax_burden",      "Annual Tax Burden — Median Path",
             "Bar chart of annual taxes paid (federal + state) on the median simulation path. "
             "The line shows the effective tax rate as a percentage of gross income."),
            ("allocation_pies", "Asset Allocation — Pre-Retirement vs. Retirement",
             "Donut charts showing the target asset allocation before and after retirement. "
             "The glide path transitions linearly between these two targets."),
            ("withdrawal_schedule", "Household Withdrawal Schedule — Median Path",
             "Chart showing median annual withdrawals. Withdrawals begin at retirement age and "
             "are inflation-adjusted from today's dollars."),
            ("success_curve", "Portfolio Success — Year-by-Year Survival Probability",
             "Filled area chart showing the fraction of simulation paths where the portfolio "
             "balance remains above zero at each year. A declining curve indicates increasing "
             "probability of portfolio depletion over time."),
            ("balance_lines", "Simulated Portfolio Balances — Percentile Lines",
             "Five discrete lines showing the 10th, 25th, 50th (median), 75th, and 90th "
             "percentile portfolio balance at each year. Provides a clear view of the "
             "distribution of outcomes across all Monte Carlo paths."),
        ]

        for key, title, caption in chart_meta:
            _heading(doc, title, 2)
            doc.add_picture(charts[key], width=Inches(6.2))
            p = doc.add_paragraph(caption)
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.italic = True
            p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            doc.add_paragraph()

        doc.add_page_break()

    def _add_data_tables(self, doc: Document) -> None:
        """Add year-by-year percentile, tax, and SS income tables."""
        _heading(doc, "3. Data Tables", 1)

        ages = self.results.ages
        pv   = self.results.portfolio_values
        tx   = self.results.taxes
        ss   = self.results.ss_income
        wd   = self.results.withdrawals

        # -- 3a: Portfolio percentiles ----------------------------------------
        _heading(doc, "3a. Portfolio Value by Percentile (every 5 years)", 2)

        step_ages = [a for a in ages if int(a) % 5 == 0]
        step_idxs = [list(ages).index(a) for a in step_ages if a in list(ages)]

        pct_rows = []
        for idx in step_idxs:
            a = int(ages[idx])
            pct_rows.append([
                f"Age {a}",
                _fmt_m(np.percentile(pv[:, idx], 10)),
                _fmt_m(np.percentile(pv[:, idx], 25)),
                _fmt_m(np.median(pv[:, idx])),
                _fmt_m(np.percentile(pv[:, idx], 75)),
                _fmt_m(np.percentile(pv[:, idx], 90)),
            ])

        _add_table(doc,
            headers=["Age", "10th", "25th", "Median", "75th", "90th"],
            rows=pct_rows,
            col_widths=[0.9, 1.1, 1.1, 1.1, 1.1, 1.1]
        )
        doc.add_paragraph()

        # -- 3b: Tax schedule -------------------------------------------------
        _heading(doc, "3b. Annual Tax and Withdrawal Schedule — Median Path (every 5 years)", 2)

        tax_rows = []
        for idx in step_idxs:
            a     = int(ages[idx])
            m_wd  = np.median(wd[:, idx])
            m_tx  = np.median(tx[:, idx])
            m_ss  = np.median(ss[:, idx])
            if hasattr(self.results, 'gross_income') and self.results.gross_income.size > 0:
                m_gross = np.median(self.results.gross_income[:, idx])
            else:
                m_gross = 0.0
            eff   = m_tx / m_gross if m_gross > 0 else 0
            tax_rows.append([
                f"Age {a}",
                _fmt_m(m_wd),
                _fmt_m(m_tx),
                _pct(eff),
                _fmt_m(m_ss),
            ])

        _add_table(doc,
            headers=["Age", "Gross Withdrawal", "Tax Paid", "Eff. Rate", "SS Income"],
            rows=tax_rows,
            col_widths=[0.9, 1.4, 1.1, 1.0, 1.2]
        )
        doc.add_paragraph()

        # -- 3b2: Withdrawal breakdown by account type ------------------------
        r = self.results
        if hasattr(r, 'trad_withdrawals') and r.trad_withdrawals.size > 0:
            _heading(doc, "3b2. Withdrawal Breakdown by Account Type — Median Path (every 5 years)", 2)
            _strategy = getattr(self.config, 'withdrawal_method', 'fixed_real').replace('_', ' ').title()
            _body(doc, f"Withdrawal strategy: {_strategy}")
            wdb_rows = []
            _cum = 0.0
            for idx in step_idxs:
                a = int(ages[idx])
                _t = np.median(r.trad_withdrawals[:, idx])
                _r = np.median(r.roth_withdrawals[:, idx])
                _x = np.median(r.taxable_withdrawals[:, idx])
                _total = np.median(wd[:, idx])
                _cum += _total
                wdb_rows.append([
                    f"Age {a}",
                    _fmt_dollar(_t),
                    _fmt_dollar(_r),
                    _fmt_dollar(_x),
                    _fmt_dollar(_total),
                    _fmt_dollar(_cum),
                ])
            _add_table(doc,
                headers=["Age", "Traditional", "Roth", "Taxable", "Annual Total", "Cumulative"],
                rows=wdb_rows,
                col_widths=[0.8, 1.0, 0.8, 0.9, 1.0, 1.0]
            )
            doc.add_paragraph()

        # -- 3c: Success analysis ---------------------------------------------
        _heading(doc, "3c. Success Analysis", 2)

        failed_mask   = ~self.results.success_mask
        n_failed      = failed_mask.sum()
        n_sims        = self.results.portfolio_values.shape[0]

        if n_failed > 0:
            # Find depletion age for each failed simulation
            depletion_ages = []
            for sim_idx in np.where(failed_mask)[0]:
                path = pv[sim_idx]
                dep  = np.argmax(path <= 0)
                if dep > 0:
                    depletion_ages.append(int(ages[dep]))
            if depletion_ages:
                dep_arr = np.array(depletion_ages)
                _add_table(doc,
                    headers=["Metric", "Value"],
                    rows=[
                        ["Simulations where portfolio depleted", f"{n_failed:,} of {n_sims:,}"],
                        ["Median depletion age",  str(int(np.median(dep_arr)))],
                        ["Earliest depletion age", str(int(np.min(dep_arr)))],
                        ["Latest depletion age",  str(int(np.max(dep_arr)))],
                    ],
                    col_widths=[4.0, 2.0]
                )
        else:
            _body(doc, f"Portfolio survived through age {self.config.life_expectancy} "
                       f"in all {n_sims:,} simulations.")

        doc.add_page_break()

    def _add_income_breakdown_table(self, doc: Document) -> None:
        """Add detailed income breakdown by age showing every income source."""
        r = self.results
        c = self.config

        _heading(doc, "Income Breakdown by Age", 2)
        _body(doc,
            "Median-path breakdown of every income source by age, with total gross "
            "income, tax paid, and effective tax rate. Deterministic income sources "
            "(SERP, pension, rental) are from config; stochastic values "
            "(tax, gross income) are median across all simulation paths.")
        doc.add_paragraph()

        from models.social_security import SocialSecurityModel
        ss_model = SocialSecurityModel()

        rows = []
        for idx, age in enumerate(r.ages):
            age_int = int(age)
            if idx > 5 and idx % 5 != 0 and idx != len(r.ages) - 1:
                continue

            cum_inf = float(np.median(
                np.cumprod(1.0 + r.inflation_rates, axis=1)[:, idx]
            )) if r.inflation_rates.size > 0 else 1.0 + 0.025 * idx
            overlays = ss_model.get_income_overlays(age_int, c, cum_inf)

            p_serp = overlays.get('serp_income', 0.0)
            pension = overlays.get('pension_income', 0.0)
            rental = overlays.get('rental_income', 0.0)
            ss = overlays.get('ss_income', 0.0)

            if hasattr(r, 'gross_income') and r.gross_income.size > 0:
                gross = float(np.median(r.gross_income[:, idx]))
            else:
                gross = 0.0

            med_tax = float(np.median(r.taxes[:, idx]))
            eff_rate = med_tax / gross if gross > 0 else 0.0

            def _v(val):
                return _fmt_dollar(val) if val > 0 else "--"

            rows.append([
                f"Age {age_int}",
                _v(p_serp), _v(pension),
                _v(rental), _v(ss),
                _fmt_dollar(gross), _fmt_dollar(med_tax), _pct(eff_rate),
            ])

        _add_table(doc,
            headers=["Age", "SERP", "Pension",
                     "Rental", "SS", "Gross", "Tax", "Eff. Rate"],
            rows=rows,
            col_widths=[0.6, 0.7, 0.7, 0.7, 0.7, 0.7, 0.7, 0.7]
        )
        doc.add_paragraph()
        doc.add_page_break()

    def _add_common_sense_check(self, doc: Document) -> None:
        """Add common sense check summary table for quick human validation."""
        r = self.results
        c = self.config

        _heading(doc, "Common Sense Check — Median Path", 2)
        _body(doc,
            "Per-year summary of key financial flows for quick validation. "
            "All values are median across simulation paths in nominal dollars. "
            "Use this table to verify that income, spending, taxes, and portfolio "
            "values are internally consistent and reasonable.")
        doc.add_paragraph()

        rows = []
        for idx, age in enumerate(r.ages):
            age_int = int(age)
            if idx > 5 and idx % 5 != 0 and idx != len(r.ages) - 1:
                continue

            if hasattr(r, 'gross_income') and r.gross_income.size > 0:
                gross = float(np.median(r.gross_income[:, idx]))
            else:
                gross = 0.0

            if hasattr(r, 'spend_amounts') and r.spend_amounts.size > 0:
                spend = float(np.median(r.spend_amounts[:, idx]))
            else:
                spend = 0.0

            tax = float(np.median(r.taxes[:, idx]))
            draw = float(np.median(r.withdrawals[:, idx]))
            portfolio = float(np.median(r.portfolio_values[:, idx]))

            rows.append([
                f"Age {age_int}",
                _fmt_dollar(gross),
                _fmt_dollar(spend) if spend > 0 else "--",
                _fmt_dollar(tax),
                _fmt_dollar(draw) if draw > 0 else "--",
                _fmt_dollar(portfolio),
            ])

        _add_table(doc,
            headers=["Age", "Total Income", "Total Spend", "Total Tax",
                     "Portfolio Draw", "Remaining Portfolio"],
            rows=rows,
            col_widths=[0.7, 1.1, 1.0, 0.9, 1.0, 1.2]
        )
        doc.add_paragraph()
        doc.add_page_break()

    def _add_rich_statistics(self, doc: Document) -> None:
        """Add rich portfolio statistics summary table (8 metrics x 5 percentiles)."""
        _heading(doc, "4. Portfolio Statistics Summary", 1)
        _body(doc,
            "Cross-simulation statistics at the 10th, 25th, 50th (median), 75th, and "
            "90th percentiles. TWRR = time-weighted rate of return. Max Drawdown measures "
            "peak-to-trough decline. SWR = safe withdrawal rate preserving portfolio to "
            "life expectancy. PWR = perpetual withdrawal rate preserving real principal."
        )
        doc.add_paragraph()

        try:
            stats = self.results.compute_rich_statistics()
        except Exception as e:
            _body(doc, f"Could not compute rich statistics: {e}")
            doc.add_page_break()
            return

        headers = ["Metric", "10th", "25th", "Median", "75th", "90th"]
        pct_keys = [10, 25, 50, 75, 90]

        # Human-readable labels for metric keys
        _METRIC_LABELS = {
            'twrr_nominal':             'TWRR (Nominal)',
            'twrr_real':                'TWRR (Real)',
            'end_balance_nominal':      'End Balance (Nominal)',
            'end_balance_real':         'End Balance (Real)',
            'max_drawdown':             'Max Drawdown',
            'max_drawdown_ex_cf':       'Max Drawdown (ex Cash Flows)',
            'safe_withdrawal_rate':     'Safe Withdrawal Rate',
            'perpetual_withdrawal_rate':'Perpetual Withdrawal Rate',
        }

        rows = []
        for label, vals in stats.items():
            # Skip scalar entries (survival_rate) -- these are
            # displayed separately, not in the percentile table.
            if not isinstance(vals, dict):
                continue

            display_label = _METRIC_LABELS.get(label, label)

            # Determine format: percentage or dollar
            is_pct = any(k in label.lower() for k in ['twrr', 'drawdown', 'rate'])
            if is_pct:
                row = [display_label] + [_pct(vals[p]) for p in pct_keys]
            else:
                row = [display_label] + [_fmt_m(vals[p]) for p in pct_keys]
            rows.append(row)

        _add_table(doc, headers=headers, rows=rows,
                   col_widths=[2.4, 1.0, 1.0, 1.0, 1.0, 1.0])

        doc.add_paragraph()
        _body(doc,
            f"Success Rate: {self.results.success_rate * 100:.1f}% -- "
            f"probability that the portfolio survives to life "
            f"expectancy (age {self.config.life_expectancy})."
        )
        doc.add_page_break()

    def _add_assumptions_log(self, doc: Document) -> None:
        """Add complete SimulationConfig parameter log for auditor / CPA review."""
        _heading(doc, "5. Assumptions Log — Auditor / CPA Reference", 1)
        _body(doc,
            "All parameters used in this simulation are documented below. "
            "Results are fully reproducible by re-running with these exact inputs. "
            "To lock the random seed for exact reproducibility, set Random Seed "
            "to a non-zero integer in Simulation Settings."
        )
        doc.add_paragraph()

        c = self.config

        sections = [
            ("Demographics", [
                ("Current Age",                  str(c.current_age)),
                ("Retirement Age",               str(c.retirement_age)),
                ("Life Expectancy",              str(c.life_expectancy)),
                ("Spouse Age",                   str(c.spouse_age) if c.spouse_age else "N/A"),
                ("Spouse Life Expectancy",       str(c.spouse_life_expectancy)),
            ]),
            ("Portfolio", [
                ("Total Portfolio Value",        _fmt_k(c.total_portfolio_value)),
                ("Traditional / 401k / IRA",     _fmt_k(c.traditional_balance)),
                ("Roth",                          _fmt_k(c.roth_balance)),
                ("Taxable Brokerage",             _fmt_k(c.taxable_balance)),
                ("Annual Contribution (Pre-Ret)", _fmt_k(c.annual_contribution)),
                ("Contribution Growth Rate",      _pct(c.contribution_growth_rate)),
                ("Expense Ratio",                 _pct(c.expense_ratio, 2)),
                ("Advisory Fee",                  _pct(c.advisory_fee, 2)),
            ]),
            ("Withdrawal Strategy", [
                ("Method",                             c.withdrawal_method),
                ("Annual Draw (Today's $)",            _fmt_k(c.annual_withdrawal_real)),
                ("Draw Start Age",                     str(getattr(c, 'withdrawal_start_age', c.retirement_age))),
                ("Withdrawal Rate",                    _pct(c.withdrawal_rate)),
                ("Guardrail Floor",                    _pct(c.withdrawal_floor)),
                ("Guardrail Ceiling",                  _pct(c.withdrawal_ceiling)),
            ]),
            ("Simulation", [
                ("Method",           c.simulation_method.upper()),
                ("Simulations",      f"{c.n_simulations:,}"),
                ("Inflation Method", c.inflation_method),
                ("Inflation Mean",   _pct(c.inflation_mean)),
                ("Inflation Std",    _pct(c.inflation_std)),
                ("HC Inflation Premium", _pct(c.healthcare_inflation_premium)),
                ("Random Seed",      str(c.random_seed) if c.random_seed else "None (non-reproducible)"),
            ]),
            ("Tax", [
                ("Filing Status",    c.filing_status.replace("_", " ").title()),
                ("State Tax Rate",   _pct(c.state_tax_rate)),
                ("Include State Tax", str(c.include_state_tax)),
            ]),
            ("Social Security", [
                ("Monthly Benefit at FRA",         f"${c.ss_monthly_benefit_at_fra:,.0f}"),
                ("Annualized Benefit at FRA",       _fmt_k(c.ss_monthly_benefit_at_fra * 12)),
                ("Full Retirement Age",             str(c.ss_fra)),
                ("Claiming Age",                    str(c.ss_claiming_age)),
                ("Include Spousal Benefit",         str(c.include_spousal_benefit)),
                ("Spouse Monthly Benefit at FRA",   f"${c.spouse_ss_monthly_benefit_at_fra:,.0f}"),
                ("Spouse Claiming Age",             str(c.spouse_ss_claiming_age)),
            ]),
            ("Other Income", [
                ("Pension (Annual Real)",          _fmt_k(c.pension_annual_real)),
                ("Pension Start Age",              str(c.pension_start_age)),
                ("Pension End Age",                str(c.pension_end_age)),
                ("Pension Tax Treatment",          "Ordinary income -- federal + state"),
                ("Rental Income (Annual Net)",     _fmt_k(c.rental_income_annual_real)),
                ("Rental Start Age",               str(c.rental_start_age)),
                ("Rental End Age",                 str(c.rental_end_age)),
                ("Rental Tax Treatment",           "Passive ordinary income -- Schedule E net"),
                ("Part-Time Income (Annual)",      _fmt_k(c.part_time_income_annual)),
                ("Part-Time Start Age",            str(c.part_time_income_start_age)),
                ("Part-Time End Age",              str(c.part_time_income_end_age)),
                ("Part-Time Tax Treatment",        "Ordinary income -- FICA applies"),
                ("-- Real / Nominal Flags --",     ""),
                ("Pension Amount Type",            "Real (inflation-adjusted)" if getattr(c, 'pension_is_real', False) else "Nominal (fixed dollar)"),
                ("Part-Time Amount Type",          "Real (inflation-adjusted)" if getattr(c, 'part_time_is_real', True) else "Nominal (fixed dollar)"),
            ]),
            ("SERP / Nonqualified Deferred Compensation (IRC 409A)",
                (lambda _rows: _rows if _rows else [("Annual Distribution", "$0")])(
                    [(f"{yr} Distribution (age {c.current_age + (yr - c.simulation_start_year)})",
                      _fmt_k(getattr(c, f'serp_{yr}', 0.0)))
                     for yr in range(2026, 2034)
                     if getattr(c, f'serp_{yr}', 0.0) > 0]
                ) + [
                ("Federal Tax Treatment",          "Ordinary income -- included in AGI"),
                ("State Tax Treatment",            "Ordinary income -- taxed per state rate"),
                ("FICA at Distribution",           "NOT applicable -- FICA paid at deferral"),
                ("Governing Code Section",          "IRC 409A"),
            ]),
            ("Spend", [
                ("Annual Spend (Today's $)",      _fmt_k(getattr(c, 'spend_annual_real', 0.0))),
                ("Spend Start Age",               str(getattr(c, 'spend_start_age', c.retirement_age))),
                ("Surplus Mode",                  getattr(c, 'spend_surplus_mode', 'ignore')),
                ("Healthcare Included",           str(c.include_healthcare_costs)),
                ("Healthcare Annual Cost",        _fmt_k(c.annual_healthcare_cost_real)),
                ("Healthcare Inflation-Adjusted", str(getattr(c, 'healthcare_is_real', True))),
                ("Healthcare Inflation Premium",  _pct(c.healthcare_inflation_premium)),
                ("Note",                          "Spend is netted against total income. "
                                                  "Shortfall drawn from portfolio. "
                                                  "Healthcare is separate with its own inflation premium."),
            ]),
            ("Asset Allocation — Pre-Retirement", [
                (a.replace("_", " ").title(), _pct(v))
                for a, v in c.pre_retirement_allocation.items()
            ]),
            ("Asset Allocation — Retirement", [
                (a.replace("_", " ").title(), _pct(v))
                for a, v in c.retirement_allocation.items()
            ]),
        ]

        for section_title, rows in sections:
            _heading(doc, section_title, 2)
            _add_table(doc,
                headers=["Parameter", "Value"],
                rows=rows,
                col_widths=[3.8, 2.5]
            )
            doc.add_paragraph()

    # -- PDF conversion -------------------------------------------------------

    def _patch_docx(self, path: str) -> None:
        """
        Patch known python-docx XML quirks in the output file.

        Fixes w:zoom missing w:percent attribute in word/settings.xml,
        which causes schema validation failure.

        Args:
            path: Path to the .docx file to patch in-place.
        """
        import tempfile
        tmp = path + ".tmp"
        with zipfile.ZipFile(path, 'r') as zin:
            with zipfile.ZipFile(tmp, 'w', compression=zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    if item.filename == "word/settings.xml":
                        xml = data.decode("utf-8")
                        xml = re.sub(
                            r'(<w:zoom\b)(?![^>]*w:percent)([^/]*/?>)',
                            r'\1 w:percent="100"\2',
                            xml
                        )
                        data = xml.encode("utf-8")
                    zout.writestr(item, data)
        os.replace(tmp, path)

    def _convert_to_pdf(self, docx_path: str) -> Optional[str]:
        """
        Convert .docx to .pdf using LibreOffice command-line interface.

        Args:
            docx_path: Absolute path to the .docx file.

        Returns:
            Absolute path to the generated .pdf, or None if conversion failed.
        """
        import sys as _sys

        output_dir = os.path.dirname(docx_path)
        _convert_args = ["--headless", "--convert-to", "pdf",
                         "--outdir", output_dir, docx_path]

        # Build candidate commands in priority order
        candidates = []

        # 1. Windows: LibreOffice in standard install locations
        if _sys.platform == 'win32':
            for env_var in ['PROGRAMFILES', 'PROGRAMFILES(X86)']:
                prog_dir = os.environ.get(env_var, '')
                if prog_dir:
                    soffice_win = os.path.join(
                        prog_dir, 'LibreOffice', 'program', 'soffice.exe')
                    if os.path.exists(soffice_win):
                        candidates.append([soffice_win] + _convert_args)

        # 2. Unix/Mac: soffice and libreoffice on PATH
        for binary in ["soffice", "libreoffice"]:
            candidates.append([binary] + _convert_args)

        for cmd in candidates:
            try:
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
                if result.returncode == 0:
                    base     = os.path.splitext(os.path.basename(docx_path))[0]
                    pdf_path = os.path.join(output_dir, f"{base}.pdf")
                    if os.path.exists(pdf_path):
                        logger.info("PDF created: %s", pdf_path)
                        return pdf_path
            except (subprocess.TimeoutExpired, FileNotFoundError, OSError):
                continue

        logger.warning("PDF conversion failed: LibreOffice not available.")
        return None
